import pyperclip
from docx import Document

def clean_text(text):
    """Elimina espacios innecesarios y asegúrate de que el texto esté en UTF-8."""
    cleaned_text = ' '.join(text.split())
    return cleaned_text

def split_text(text, max_length=4096):
    """Divide el texto en fragmentos de un tamaño máximo especificado."""
    fragments = []
    while len(text) > max_length:
        split_point = text.rfind(' ', 0, max_length)
        if split_point == -1:
            split_point = max_length
        fragments.append(text[:split_point])
        text = text[split_point:].strip()
    fragments.append(text)
    return fragments

def save_fragments_to_word(fragments, filename='output.docx'):
    """Guarda todos los fragmentos en un archivo de Word."""
    doc = Document()
    for i, fragment in enumerate(fragments):
        doc.add_heading(f'Fragment {i+1}', level=1)
        doc.add_paragraph(fragment)
        doc.add_paragraph('-' * 40)
    doc.save(filename)
    print(f"Saved all fragments to {filename}")

def main():
    input_text = pyperclip.paste()  # Lee el texto del portapapeles
    cleaned_text = clean_text(input_text)
    fragments = split_text(cleaned_text)
    
    save_fragments_to_word(fragments, 'output.docx')

if __name__ == "__main__":
    main()
